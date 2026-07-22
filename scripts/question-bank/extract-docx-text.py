import sys
from pathlib import Path

from docx import Document


def main():
    if len(sys.argv) != 3:
        raise SystemExit("usage: extract-docx-text.py INPUT.docx OUTPUT.txt")
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    document = Document(source)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append("\t".join(cell.text for cell in row.cells))
    output.write_text("\n".join(paragraphs) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
